"""
Procedure CRUD service + docx extraction + the upload endpoint (incl. 10 MB cap).
"""

from io import BytesIO

import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.exceptions import ConflictException, ValidationException
from app.procedures.constants import ProcedureStatus
from app.procedures.controllers.procedure_controller import router as procedures_router
from app.procedures.repositories import (
    ProcedureRepository,
    QuizAttemptRepository,
    QuizQuestionRepository,
)
from app.procedures.services.docx_service import (
    _sanitize_procedure_html,
    extract_html_from_docx,
    extract_text_from_docx,
)
from app.procedures.services.procedure_service import ProcedureService
from app.repositories.system_settings_repository import SystemSettingsRepository
from app.repositories.user_repository import UserRepository
from app.services.settings_service import SettingsService


class _FakePublisher:
    def __init__(self):
        self.calls = []

    async def broadcast(self, recipients, title, procedure_id):
        self.calls.append((list(recipients), title, procedure_id))
        return {"sent": len(recipients), "skipped": 0, "total": len(recipients)}


def _svc(db_session, publisher=None) -> ProcedureService:
    return ProcedureService(
        ProcedureRepository(db_session),
        QuizQuestionRepository(db_session),
        QuizAttemptRepository(db_session),
        UserRepository(db_session),
        SettingsService(SystemSettingsRepository(db_session)),
        publisher or _FakePublisher(),
    )


def _docx_bytes(paragraphs):
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── CRUD service ────────────────────────────────────────────────────────────


async def test_create_is_draft(db_session):
    svc = _svc(db_session)
    proc = await svc.create("נהל אבטחה", "תוכן הנהל", source_filename="a.docx")
    assert proc.status == ProcedureStatus.DRAFT
    assert proc.source_filename == "a.docx"


async def test_list_all_includes_question_counts(db_session):
    svc = _svc(db_session)
    proc = await svc.create("נהל", "תוכן")
    from app.procedures.constants import QuestionSource
    from app.procedures.models import QuizQuestion

    db_session.add(
        QuizQuestion(
            procedure_id=proc.id, text="q1", options=["a", "b"],
            correct_index=0, display_order=0, source=QuestionSource.AI,
        )
    )
    db_session.add(
        QuizQuestion(
            procedure_id=proc.id, text="q2", options=["a", "b"],
            correct_index=0, display_order=1, is_active=False, source=QuestionSource.AI,
        )
    )
    await db_session.commit()
    rows = await svc.list_all()
    assert rows[0]["active_questions"] == 1
    assert rows[0]["total_questions"] == 2


async def test_get_with_questions_and_404(db_session):
    svc = _svc(db_session)
    proc = await svc.create("נהל", "תוכן")
    got = await svc.get(proc.id)
    assert got.questions == []
    import uuid

    with pytest.raises(Exception):
        await svc.get(uuid.uuid4())


async def test_update_only_while_draft(db_session):
    svc = _svc(db_session)
    proc = await svc.create("נהל", "תוכן")
    updated = await svc.update(proc.id, title="נהל חדש", body_text=None)
    assert updated.title == "נהל חדש"
    # Force-publish to test the guard.
    proc = await svc._procedures.get_by_id(proc.id)
    await svc._procedures.update(proc.id, status=ProcedureStatus.PUBLISHED)
    with pytest.raises(ConflictException):
        await svc.update(proc.id, title="אסור", body_text=None)


async def test_archive_sets_archived(db_session):
    svc = _svc(db_session)
    proc = await svc.create("נהל", "תוכן")
    archived = await svc.archive(proc.id)
    assert archived.status == ProcedureStatus.ARCHIVED
    # idempotent
    again = await svc.archive(proc.id)
    assert again.status == ProcedureStatus.ARCHIVED


# ── docx extraction ─────────────────────────────────────────────────────────


def test_extract_text_from_valid_docx():
    data = _docx_bytes(["סעיף אחד", "סעיף שני"])
    text = extract_text_from_docx(data)
    assert "סעיף אחד" in text
    assert "סעיף שני" in text
    assert "\n\n" in text


def test_extract_text_invalid_docx_raises_validation():
    with pytest.raises(ValidationException):
        extract_text_from_docx(b"not a real docx file content")


def test_extract_text_includes_tables():
    from docx import Document

    doc = Document()
    doc.add_paragraph("הקדמה")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "עמדה"
    table.rows[0].cells[1].text = "שעה"
    buf = BytesIO()
    doc.save(buf)
    text = extract_text_from_docx(buf.getvalue())
    assert "הקדמה" in text
    assert "עמדה" in text
    assert "שעה" in text


# ── docx extraction: bold markers ───────────────────────────────────────────


def _save_docx(doc) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_bold_run_becomes_markers():
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("מלל רגיל ")
    bold = p.add_run("טקסט מודגש")
    bold.bold = True
    text = extract_text_from_docx(_save_docx(doc))
    # the bold run is wrapped in *…*, the plain run is not
    assert text == "מלל רגיל *טקסט מודגש*"


def test_docx_adjacent_bold_runs_merge_into_one_span():
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("אב")
    r1.bold = True
    r2 = p.add_run("גד")
    r2.bold = True
    text = extract_text_from_docx(_save_docx(doc))
    # merged into a single *…* span — never *אב**גד*
    assert text == "*אבגד*"
    assert "**" not in text


def test_docx_heading_style_marked_bold():
    from docx import Document

    doc = Document()
    doc.add_heading("כותרת חשובה", level=1)
    text = extract_text_from_docx(_save_docx(doc))
    # a Heading-style paragraph is marked even if its runs don't set bold
    assert text == "*כותרת חשובה*"


def test_docx_highlighted_run_marked_bold():
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("רגיל ")
    hl = p.add_run("מודגש בצבע")
    hl.font.highlight_color = WD_COLOR_INDEX.YELLOW
    text = extract_text_from_docx(_save_docx(doc))
    assert "*מודגש בצבע*" in text
    assert text.startswith("רגיל ")  # the plain run is not wrapped


def test_docx_plain_text_has_no_markers():
    from docx import Document

    doc = Document()
    doc.add_paragraph("פסקה רגילה לחלוטין")
    text = extract_text_from_docx(_save_docx(doc))
    assert text == "פסקה רגילה לחלוטין"
    assert "*" not in text


def test_docx_whitespace_only_bold_run_gets_no_markers():
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("לפני ")
    r2 = p.add_run(" ")  # whitespace-only bold, surrounded by non-bold runs
    r2.bold = True
    r3 = p.add_run("אחרי")
    text = extract_text_from_docx(_save_docx(doc))
    # a whitespace-only bold run emits no markers
    assert "*" not in text
    assert "לפני" in text and "אחרי" in text


def test_docx_markers_never_span_paragraphs():
    from docx import Document

    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run("התחלה").bold = True
    p2 = doc.add_paragraph()
    p2.add_run("סוף").bold = True
    text = extract_text_from_docx(_save_docx(doc))
    # two independent *…* spans separated by a paragraph break, not one span
    assert text == "*התחלה*\n\n*סוף*"


# ── Upload endpoint (10 MB cap + happy path) ─────────────────────────────────


def _client():
    from app.dependencies import require_admin_role

    app = FastAPI()
    app.include_router(procedures_router)
    app.dependency_overrides[require_admin_role] = lambda: None
    return TestClient(app)


def test_upload_docx_returns_extracted_text():
    data = _docx_bytes(["נהל כניסה לאתר"])
    client = _client()
    res = client.post(
        "/admin/procedures/upload",
        files={"file": ("proc.docx", data, "application/vnd.openxmlformats")},
        data={"title": "נהל"},
    )
    assert res.status_code == 200
    body = res.json()
    assert "נהל כניסה לאתר" in body["text"]
    assert body["source_filename"] == "proc.docx"
    assert body["char_count"] == len(body["text"])


def test_upload_docx_10mb_cap(monkeypatch):
    """A file over the cap is rejected (413) before extraction."""
    from app.procedures.controllers import procedure_controller as ctrl

    monkeypatch.setattr(ctrl, "MAX_DOCX_BYTES", 64)
    data = _docx_bytes(["x" * 200])  # bigger than the patched 64-byte cap
    client = _client()
    res = client.post(
        "/admin/procedures/upload",
        files={"file": ("big.docx", data, "application/vnd.openxmlformats")},
        data={"title": "נהל"},
    )
    assert res.status_code == 413


@pytest.mark.asyncio
async def test_create_endpoint_serializes_fresh_procedure(db_session):
    """Regression: the create endpoint must not lazy-load `questions` off the
    freshly-created instance (MissingGreenlet on an async session) — it re-fetches
    with the relationship eager-loaded before serializing."""
    from app.procedures.controllers.procedure_controller import create_procedure
    from app.procedures.schemas import ProcedureCreate

    out = await create_procedure(
        ProcedureCreate(title="נוהל חדש", body_text="תוכן הנוהל"),
        service=_svc(db_session),
    )
    assert out.title == "נוהל חדש"
    assert out.status == "draft"
    assert out.questions == []


@pytest.mark.asyncio
async def test_delete_procedure_removes_it_and_its_questions(db_session):
    """Hard delete: the procedure and its question bank are gone (ORM cascade);
    a later get raises 404."""
    from app.exceptions import UserNotFoundException
    from app.procedures.constants import QuestionSource

    svc = _svc(db_session)
    proc = await svc.create("נוהל למחיקה", "תוכן")
    await QuizQuestionRepository(db_session).create(
        procedure_id=proc.id,
        text="שאלה?",
        options=["א", "ב"],
        correct_index=0,
        display_order=0,
        is_active=True,
        source=QuestionSource.AI,
    )

    await svc.delete(proc.id)

    assert await QuizQuestionRepository(db_session).count_all(proc.id) == 0
    with pytest.raises(UserNotFoundException):
        await svc.get(proc.id)


@pytest.mark.asyncio
async def test_delete_unknown_procedure_404(db_session):
    import uuid as _uuid

    from app.exceptions import UserNotFoundException

    with pytest.raises(UserNotFoundException):
        await _svc(db_session).delete(_uuid.uuid4())


@pytest.mark.asyncio
async def test_list_all_flags_has_ai_questions(db_session):
    """list_all exposes has_ai_questions: True only when an AI question exists
    (manual-only banks keep the generate button available in the UI)."""
    from app.procedures.constants import QuestionSource

    svc = _svc(db_session)
    ai_proc = await svc.create("נוהל עם בנק AI", "תוכן")
    manual_proc = await svc.create("נוהל ידני", "תוכן")
    repo = QuizQuestionRepository(db_session)
    await repo.create(
        procedure_id=ai_proc.id, text="ש?", options=["א", "ב"], correct_index=0,
        display_order=0, is_active=True, source=QuestionSource.AI,
    )
    await repo.create(
        procedure_id=manual_proc.id, text="ש?", options=["א", "ב"], correct_index=0,
        display_order=0, is_active=True, source=QuestionSource.MANUAL,
    )

    rows = {r["title"]: r["has_ai_questions"] for r in await svc.list_all()}
    assert rows["נוהל עם בנק AI"] is True
    assert rows["נוהל ידני"] is False


# ── docx → sanitized HTML extraction (body_html) ─────────────────────────────


def _rich_docx_bytes() -> bytes:
    """A docx with a heading, a bold run, a bullet list, and a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("כותרת ראשית", level=1)
    p = doc.add_paragraph()
    p.add_run("טקסט רגיל ")
    bold = p.add_run("מילה מודגשת")
    bold.bold = True
    doc.add_paragraph("סעיף ראשון", style="List Bullet")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "עמודה א"
    table.rows[0].cells[1].text = "עמודה ב"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_html_from_rich_docx_has_structure():
    """mammoth converts heading/bold/list/table; sanitize keeps the tags. The
    plain-text extraction is unaffected by the new HTML path."""
    data = _rich_docx_bytes()
    html = extract_html_from_docx(data)
    assert html is not None
    assert "<h1>" in html and "כותרת ראשית" in html
    assert "<strong>" in html and "מילה מודגשת" in html
    assert "<li>" in html and "סעיף ראשון" in html
    assert "<td>" in html and "עמודה א" in html and "עמודה ב" in html
    # text extraction still works on the same bytes (unchanged behavior)
    text = extract_text_from_docx(data)
    assert "כותרת ראשית" in text


def test_extract_html_strips_xss_payloads():
    """Sanitizer drops <script> (content too), onclick attrs, and <img>. [EDGE D5]"""
    crafted = (
        '<p>תוכן תקין</p>'
        '<script>alert("xss")</script>'
        '<p onclick="steal()">בוצע קליק</p>'
        '<img src="x" onerror="alert(1)">'
    )
    sanitized = _sanitize_procedure_html(crafted)
    assert "<script>" not in sanitized
    assert "alert" not in sanitized  # script content removed, not just the tag
    assert "onclick" not in sanitized
    assert "<img" not in sanitized
    # the legitimate text survives
    assert "תוכן תקין" in sanitized
    assert "בוצע קליק" in sanitized


def test_extract_html_corrupt_docx_returns_none_without_raising():
    """A non-docx payload: text extraction still raises (validity gate), HTML
    conversion returns None and never raises. [EDGE D4]"""
    bad = b"not a real docx file content"
    with pytest.raises(ValidationException):
        extract_text_from_docx(bad)
    assert extract_html_from_docx(bad) is None


def test_extract_html_empty_result_returns_none():
    """A docx whose conversion yields only whitespace → None (fallback path)."""
    # _sanitize_procedure_html on whitespace-only input → None at the extractor
    # boundary (extract_html_from_docx treats empty/whitespace as None).
    assert extract_html_from_docx(b"") is None


@pytest.mark.asyncio
async def test_create_persists_and_returns_body_html(db_session):
    svc = _svc(db_session)
    proc = await svc.create(
        "נהל עם HTML", "תוכן", body_html="<p>סקירה</p>", source_filename="x.docx"
    )
    assert proc.body_html == "<p>סקירה</p>"
    fetched = await svc.get(proc.id)
    assert fetched.body_html == "<p>סקירה</p>"


@pytest.mark.asyncio
async def test_update_without_body_html_preserves_existing(db_session):
    """Omitting body_html on update must NOT clear the existing snapshot. [EDGE D3]"""
    svc = _svc(db_session)
    proc = await svc.create("נהל", "תוכן", body_html="<p>ישן</p>")
    await svc.update(proc.id, title="כותרת חדשה", body_text="תוכן חדש")
    fetched = await svc.get(proc.id)
    assert fetched.title == "כותרת חדשה"
    assert fetched.body_text == "תוכן חדש"
    # snapshot untouched by a text-only edit
    assert fetched.body_html == "<p>ישן</p>"


@pytest.mark.asyncio
async def test_update_with_body_html_replaces_existing(db_session):
    svc = _svc(db_session)
    proc = await svc.create("נהל", "תוכן", body_html="<p>ישן</p>")
    await svc.update(proc.id, title=None, body_text=None, body_html="<p>חדש</p>")
    fetched = await svc.get(proc.id)
    assert fetched.body_html == "<p>חדש</p>"


def test_upload_docx_returns_body_html():
    """The upload endpoint response carries the sanitized body_html snapshot."""
    data = _rich_docx_bytes()
    client = _client()
    res = client.post(
        "/admin/procedures/upload",
        files={"file": ("proc.docx", data, "application/vnd.openxmlformats")},
        data={"title": "נהל"},
    )
    assert res.status_code == 200
    body = res.json()
    assert body["body_html"] is not None
    assert "<h1>" in body["body_html"]
    assert "<strong>" in body["body_html"]
