"""
docx text extraction for procedure uploads.

Extracts the plain text of an uploaded Word document so the admin can review it
before saving (tables are flattened to text; images/headers/footers are out of
scope — accepted limitation, the admin edits the text before saving). Bold runs
(plus Heading-style paragraphs and highlighted runs) are preserved as lightweight
``*…*`` markers so the guard WebApp fallback renderer (``body_text`` → ``<strong>``)
and the AI question generator can treat them as emphasis; markers never span a
paragraph.

``extract_html_from_docx`` produces the rich-HTML snapshot consumed by the guard
WebApp reading page: mammoth converts the docx (headings/lists/tables/bold) and
the result is sanitized server-side with nh3 (allowlist of structural tags;
attributes stripped except colspan/rowspan; images dropped). It NEVER raises —
the plain-text extractor stays the sole upload validity gate, so an exotic or
image-only docx that fails HTML conversion still uploads exactly as today (the
page then falls back to rendering ``body_text``).
"""

import io

import nh3

from app.exceptions import ValidationException

# Hard upload cap. Enforced in the controller before extraction so a huge upload
# is rejected without being fully buffered.
MAX_DOCX_BYTES = 10 * 1024 * 1024  # 10 MB

# nh3 allowlist for the sanitized procedure HTML. Structural/formatting tags only
# — no links, no images, no scripts. The frontend injects this field raw via
# dangerouslySetInnerHTML, so it MUST be sanitized here first (see [EDGE D5]).
_HTML_ALLOWED_TAGS = {
    "p", "h1", "h2", "h3", "h4", "h5", "h6",
    "strong", "b", "em", "i", "u", "s", "br",
    "blockquote", "ol", "ul", "li",
    "table", "thead", "tbody", "tr", "th", "td",
}
# Only colspan/rowspan survive on table cells; every other attribute is stripped.
_HTML_ALLOWED_ATTRS = {
    "td": {"colspan", "rowspan"},
    "th": {"colspan", "rowspan"},
}


def _sanitize_procedure_html(html: str) -> str:
    """Sanitize mammoth's HTML to the procedure allowlist (XSS hardening).

    ``clean_content_tags`` drops ``<script>``/``<style>`` *content* entirely
    (a stripped-but-not-content-cleaned script would leave its body as visible
    text). Tags/attributes outside the allowlist are removed by nh3.
    """
    return nh3.clean(
        html,
        tags=_HTML_ALLOWED_TAGS,
        attributes=_HTML_ALLOWED_ATTRS,
        clean_content_tags={"script", "style"},
    )


def extract_html_from_docx(data: bytes) -> str | None:
    """Convert a .docx byte string to sanitized HTML for the guard WebApp page.

    mammoth's ``convert_to_html`` runs on the same bytes; images are ignored via
    a no-op ``convert_image`` (images/headers/footers stay out of scope, matching
    the text extractor's documented limitation), and the output is sanitized to
    the procedure allowlist.

    Returns the sanitized HTML, or ``None`` when conversion is impossible or
    yields nothing meaningful (exotic docx, image-only document). NEVER raises —
    the plain-text extractor (``extract_text_from_docx``) is the upload validity
    gate, so a failed/empty HTML conversion leaves ``body_html`` NULL and the
    page falls back to ``body_text`` (see [EDGE D4]).
    """
    try:
        import mammoth  # lazy: keeps module import docx-light
    except ImportError:  # pragma: no cover - dep is in requirements
        return None

    try:
        result = mammoth.convert_to_html(io.BytesIO(data), convert_image=lambda _img: [])
    except Exception:  # noqa: BLE001 — never raise (best-effort HTML snapshot)
        return None

    try:
        sanitized = _sanitize_procedure_html(result.value or "")
    except Exception:  # noqa: BLE001 — defensive: nh3 should not fail
        return None
    if not sanitized or not sanitized.strip():
        return None
    return sanitized



def extract_text_from_docx(data: bytes) -> str:
    """Extract concatenated paragraph + table text from a .docx byte string.

    Raises ValidationException if the bytes are not a valid Word document (the
    admin uploaded a renamed non-docx file).
    """
    try:
        from docx import Document  # lazy: keeps module import docx-light
    except ImportError as exc:  # pragma: no cover - dep is in requirements
        raise RuntimeError("python-docx is not installed") from exc

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValidationException("הקובץ אינו מסמך Word תקין (docx)") from exc

    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        text = _paragraph_to_marked_text(paragraph).strip()
        if text:
            chunks.append(text)
    # Flatten tables (rows/cells) so their text isn't silently lost. Tables do
    # not carry bold markers (cell.text drops run formatting) — accepted, same
    # as images/headers/footers being out of scope.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n\n".join(chunks)


def _paragraph_to_marked_text(paragraph) -> str:
    """Render one paragraph's runs as text, wrapping bold spans in ``*markers*``.

    A run counts as bold when its ``bold`` flag is set, it carries a highlight
    color, or the paragraph itself is a Heading style. Adjacent bold runs merge
    into a single ``*…*`` span (never ``*a**b*``); whitespace-only bold runs get
    no markers; markers never span paragraphs (each paragraph is independent,
    so the Telegram conversion can stay balanced per line).
    """
    try:
        style_name = paragraph.style.name or ""
    except Exception:  # pragma: no cover - defensive: malformed style refs
        style_name = ""
    is_heading = style_name.startswith("Heading")

    # Coalesce adjacent runs that share the same bold flag into one segment so
    # a bold word split across several runs becomes a single *…* span.
    segments: list[tuple[bool, str]] = []  # (is_bold, text)
    for run in paragraph.runs:
        text = run.text or ""
        if text == "":
            continue
        is_bold = (
            bool(run.bold)
            or getattr(run.font, "highlight_color", None) is not None
            or is_heading
        )
        if segments and segments[-1][0] == is_bold:
            prev_bold, prev_text = segments[-1]
            segments[-1] = (prev_bold, prev_text + text)
        else:
            segments.append((is_bold, text))

    out: list[str] = []
    for is_bold, text in segments:
        if is_bold and text.strip():
            out.append(f"*{text}*")
        else:
            out.append(text)
    return "".join(out)
